import logging
from io import BytesIO
from html import escape

from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class ResumeExporter:
    """
    Exports a Resume model instance to PDF or DOCX format.

    The resume.content JSON is expected to have keys such as:
        - personal: dict with name, email, phone, address, linkedin, website, summary
        - education: list of dicts with institution, degree, field, start_date, end_date, gpa, description
        - experience: list of dicts with company, position, start_date, end_date, location, description, achievements
        - skills: list of strings, or dict mapping category names to lists of strings
        - projects: list of dicts with name, description, technologies, url
        - certifications: list of dicts with name, issuer, date, url

    All sections are optional and handled gracefully when absent.
    """

    def __init__(self, resume):
        self.resume = resume

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def export_pdf(self) -> bytes:
        """Export resume as PDF using xhtml2pdf.

        Builds HTML from resume content JSON, applies template CSS if available,
        then converts to PDF. Returns PDF bytes.
        """
        html = self._build_html()
        buffer = BytesIO()

        pisa_status = pisa.CreatePDF(html, dest=buffer, encoding="utf-8")

        if pisa_status.err:
            logger.error(
                "PDF generation failed for resume %s: %s",
                self.resume.id,
                pisa_status.err,
            )
            raise RuntimeError("Failed to generate PDF for the resume.")

        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def export_docx(self) -> bytes:
        """Export resume as DOCX using python-docx.

        Builds a Word document from resume content JSON.
        Returns DOCX bytes.
        """
        doc = Document()
        content = self.resume.content or {}

        # ---- Page margins ----
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # ---- Personal / Header ----
        personal = content.get("personal", {})
        name = personal.get("name", self.resume.title or "Untitled Resume")

        name_para = doc.add_heading(name, level=0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in name_para.runs:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Contact line
        contact_parts = []
        if personal.get("email"):
            contact_parts.append(personal["email"])
        if personal.get("phone"):
            contact_parts.append(personal["phone"])
        if personal.get("address"):
            contact_parts.append(personal["address"])
        if personal.get("linkedin"):
            contact_parts.append(personal["linkedin"])
        if personal.get("website"):
            contact_parts.append(personal["website"])

        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Summary / Objective
        summary = personal.get("summary", "")
        if summary:
            doc.add_heading("Professional Summary", level=2)
            summary_para = doc.add_paragraph(summary)
            for run in summary_para.runs:
                run.font.size = Pt(10)

        # ---- Experience ----
        experience = content.get("experience", [])
        if experience:
            doc.add_heading("Work Experience", level=2)
            for exp in experience:
                # Position / Company line
                position = exp.get("position", "")
                company = exp.get("company", "")
                header_text = " — ".join(filter(None, [position, company]))
                if header_text:
                    exp_heading = doc.add_paragraph()
                    run = exp_heading.add_run(header_text)
                    run.bold = True
                    run.font.size = Pt(11)

                # Date / Location line
                date_parts = []
                start = exp.get("start_date", "")
                end = exp.get("end_date", "Present")
                if start:
                    date_parts.append(f"{start} - {end}")
                location = exp.get("location", "")
                if location:
                    date_parts.append(location)
                if date_parts:
                    meta_para = doc.add_paragraph(" | ".join(date_parts))
                    for run in meta_para.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                # Description
                description = exp.get("description", "")
                if description:
                    desc_para = doc.add_paragraph(description)
                    for run in desc_para.runs:
                        run.font.size = Pt(10)

                # Achievements as bullet points
                achievements = exp.get("achievements", [])
                if isinstance(achievements, list):
                    for achievement in achievements:
                        bullet = doc.add_paragraph(str(achievement), style="List Bullet")
                        for run in bullet.runs:
                            run.font.size = Pt(10)

        # ---- Education ----
        education = content.get("education", [])
        if education:
            doc.add_heading("Education", level=2)
            for edu in education:
                degree = edu.get("degree", "")
                field = edu.get("field", "")
                institution = edu.get("institution", "")
                degree_text = " in ".join(filter(None, [degree, field]))
                header_text = " — ".join(filter(None, [degree_text, institution]))
                if header_text:
                    edu_heading = doc.add_paragraph()
                    run = edu_heading.add_run(header_text)
                    run.bold = True
                    run.font.size = Pt(11)

                date_parts = []
                start = edu.get("start_date", "")
                end = edu.get("end_date", "")
                if start or end:
                    date_parts.append(f"{start} - {end}".strip(" -"))
                gpa = edu.get("gpa", "")
                if gpa:
                    date_parts.append(f"GPA: {gpa}")
                if date_parts:
                    meta_para = doc.add_paragraph(" | ".join(date_parts))
                    for run in meta_para.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                description = edu.get("description", "")
                if description:
                    desc_para = doc.add_paragraph(description)
                    for run in desc_para.runs:
                        run.font.size = Pt(10)

        # ---- Skills ----
        skills = content.get("skills", None)
        if skills:
            doc.add_heading("Skills", level=2)
            if isinstance(skills, list):
                skills_para = doc.add_paragraph(", ".join(str(s) for s in skills))
                for run in skills_para.runs:
                    run.font.size = Pt(10)
            elif isinstance(skills, dict):
                for category, items in skills.items():
                    cat_para = doc.add_paragraph()
                    cat_run = cat_para.add_run(f"{category}: ")
                    cat_run.bold = True
                    cat_run.font.size = Pt(10)
                    if isinstance(items, list):
                        items_run = cat_para.add_run(", ".join(str(s) for s in items))
                    else:
                        items_run = cat_para.add_run(str(items))
                    items_run.font.size = Pt(10)

        # ---- Projects ----
        projects = content.get("projects", [])
        if projects:
            doc.add_heading("Projects", level=2)
            for proj in projects:
                proj_name = proj.get("name", "")
                if proj_name:
                    proj_heading = doc.add_paragraph()
                    run = proj_heading.add_run(proj_name)
                    run.bold = True
                    run.font.size = Pt(11)

                description = proj.get("description", "")
                if description:
                    desc_para = doc.add_paragraph(description)
                    for run in desc_para.runs:
                        run.font.size = Pt(10)

                technologies = proj.get("technologies", "")
                if technologies:
                    if isinstance(technologies, list):
                        technologies = ", ".join(str(t) for t in technologies)
                    tech_para = doc.add_paragraph()
                    label_run = tech_para.add_run("Technologies: ")
                    label_run.bold = True
                    label_run.font.size = Pt(10)
                    tech_run = tech_para.add_run(str(technologies))
                    tech_run.font.size = Pt(10)

                url = proj.get("url", "")
                if url:
                    url_para = doc.add_paragraph()
                    label_run = url_para.add_run("URL: ")
                    label_run.bold = True
                    label_run.font.size = Pt(10)
                    url_run = url_para.add_run(str(url))
                    url_run.font.size = Pt(10)

        # ---- Certifications ----
        certifications = content.get("certifications", [])
        if certifications:
            doc.add_heading("Certifications", level=2)
            for cert in certifications:
                cert_name = cert.get("name", "")
                issuer = cert.get("issuer", "")
                cert_date = cert.get("date", "")

                header_text = cert_name
                if issuer:
                    header_text = f"{cert_name} — {issuer}" if cert_name else issuer
                if header_text:
                    cert_heading = doc.add_paragraph()
                    run = cert_heading.add_run(header_text)
                    run.bold = True
                    run.font.size = Pt(11)

                meta_parts = []
                if cert_date:
                    meta_parts.append(cert_date)
                cert_url = cert.get("url", "")
                if cert_url:
                    meta_parts.append(cert_url)
                if meta_parts:
                    meta_para = doc.add_paragraph(" | ".join(meta_parts))
                    for run in meta_para.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        # ---- Serialize to bytes ----
        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _build_html(self) -> str:
        """Build a complete HTML document string from resume content.

        If the resume has an associated template with html_structure and
        css_styles, those are used. Otherwise a clean default layout is
        generated.
        """
        content = self.resume.content or {}
        template = getattr(self.resume, "template", None)

        # Determine CSS
        if template and getattr(template, "css_styles", None):
            css = template.css_styles
        else:
            css = self._get_default_css()

        # Determine body HTML
        if template and getattr(template, "html_structure", None):
            body_html = self._render_template_html(template.html_structure, content)
        else:
            body_html = self._build_default_body(content)

        html = (
            "<!DOCTYPE html>\n"
            "<html>\n"
            "<head>\n"
            '  <meta charset="utf-8">\n'
            "  <style>\n"
            f"    {css}\n"
            "  </style>\n"
            "</head>\n"
            "<body>\n"
            f"  {body_html}\n"
            "</body>\n"
            "</html>"
        )
        return html

    def _render_template_html(self, html_structure: str, content: dict) -> str:
        """Render the template's html_structure by replacing placeholders with
        content values.

        Supports simple ``{{key}}`` and ``{{section.field}}`` placeholders.
        Unresolved placeholders are replaced with empty strings.
        """
        import re

        rendered = html_structure

        # Flatten personal fields for simple {{name}}, {{email}}, etc.
        personal = content.get("personal", {})
        for key, value in personal.items():
            placeholder = "{{" + key + "}}"
            rendered = rendered.replace(placeholder, escape(str(value)))

        # Replace {{section.field}} patterns
        pattern = re.compile(r"\{\{(\w+)\.(\w+)\}\}")
        def replacer(match):
            section_key = match.group(1)
            field_key = match.group(2)
            section_data = content.get(section_key, {})
            if isinstance(section_data, dict):
                return escape(str(section_data.get(field_key, "")))
            return ""

        rendered = pattern.sub(replacer, rendered)

        # Inject rendered section blocks for list-based sections
        for section_key in ("education", "experience", "skills", "projects", "certifications"):
            section_data = content.get(section_key)
            if section_data is None:
                continue

            section_html = self._build_section_html(section_key, section_data)
            # Replace a block placeholder like {{education_section}}
            block_placeholder = "{{" + section_key + "_section}}"
            rendered = rendered.replace(block_placeholder, section_html)

        # Remove any remaining unresolved placeholders
        rendered = re.sub(r"\{\{[^}]+\}\}", "", rendered)

        return rendered

    def _build_default_body(self, content: dict) -> str:
        """Build the full body HTML using the default layout."""
        parts = []
        personal = content.get("personal", {})

        # Header with name and contact info
        name = escape(personal.get("name", self.resume.title or ""))
        parts.append(f'<div class="header"><h1>{name}</h1>')

        contact_items = []
        if personal.get("email"):
            contact_items.append(escape(personal["email"]))
        if personal.get("phone"):
            contact_items.append(escape(personal["phone"]))
        if personal.get("address"):
            contact_items.append(escape(personal["address"]))
        if personal.get("linkedin"):
            contact_items.append(escape(personal["linkedin"]))
        if personal.get("website"):
            contact_items.append(escape(personal["website"]))

        if contact_items:
            parts.append(
                '<p class="contact">' + " &bull; ".join(contact_items) + "</p>"
            )
        parts.append("</div>")  # close .header

        # Summary
        summary = personal.get("summary", "")
        if summary:
            parts.append('<div class="section">')
            parts.append('<h2 class="section-title">Professional Summary</h2>')
            parts.append(f"<p>{escape(summary)}</p>")
            parts.append("</div>")

        # Experience
        experience = content.get("experience", [])
        if experience:
            parts.append(self._build_section_html("experience", experience))

        # Education
        education = content.get("education", [])
        if education:
            parts.append(self._build_section_html("education", education))

        # Skills
        skills = content.get("skills")
        if skills:
            parts.append(self._build_section_html("skills", skills))

        # Projects
        projects = content.get("projects", [])
        if projects:
            parts.append(self._build_section_html("projects", projects))

        # Certifications
        certifications = content.get("certifications", [])
        if certifications:
            parts.append(self._build_section_html("certifications", certifications))

        return "\n".join(parts)

    def _build_section_html(self, section_key: str, data) -> str:
        """Return an HTML block for a given resume section."""
        section_titles = {
            "experience": "Work Experience",
            "education": "Education",
            "skills": "Skills",
            "projects": "Projects",
            "certifications": "Certifications",
        }
        title = section_titles.get(section_key, section_key.title())
        lines = [
            '<div class="section">',
            f'<h2 class="section-title">{escape(title)}</h2>',
        ]

        if section_key == "experience":
            for exp in data if isinstance(data, list) else []:
                lines.append('<div class="entry">')
                position = escape(exp.get("position", ""))
                company = escape(exp.get("company", ""))
                header = " &mdash; ".join(filter(None, [position, company]))
                if header:
                    lines.append(f'<h3 class="entry-title">{header}</h3>')

                start = escape(exp.get("start_date", ""))
                end = escape(exp.get("end_date", "Present"))
                location = escape(exp.get("location", ""))
                meta_parts = []
                if start:
                    meta_parts.append(f"{start} - {end}")
                if location:
                    meta_parts.append(location)
                if meta_parts:
                    lines.append(
                        '<p class="entry-meta">' + " | ".join(meta_parts) + "</p>"
                    )

                description = exp.get("description", "")
                if description:
                    lines.append(f"<p>{escape(description)}</p>")

                achievements = exp.get("achievements", [])
                if isinstance(achievements, list) and achievements:
                    lines.append("<ul>")
                    for ach in achievements:
                        lines.append(f"  <li>{escape(str(ach))}</li>")
                    lines.append("</ul>")
                lines.append("</div>")

        elif section_key == "education":
            for edu in data if isinstance(data, list) else []:
                lines.append('<div class="entry">')
                degree = escape(edu.get("degree", ""))
                field = escape(edu.get("field", ""))
                institution = escape(edu.get("institution", ""))
                degree_text = " in ".join(filter(None, [degree, field]))
                header = " &mdash; ".join(filter(None, [degree_text, institution]))
                if header:
                    lines.append(f'<h3 class="entry-title">{header}</h3>')

                meta_parts = []
                start = escape(edu.get("start_date", ""))
                end = escape(edu.get("end_date", ""))
                if start or end:
                    meta_parts.append(f"{start} - {end}".strip(" -"))
                gpa = edu.get("gpa", "")
                if gpa:
                    meta_parts.append(f"GPA: {escape(str(gpa))}")
                if meta_parts:
                    lines.append(
                        '<p class="entry-meta">' + " | ".join(meta_parts) + "</p>"
                    )

                description = edu.get("description", "")
                if description:
                    lines.append(f"<p>{escape(description)}</p>")
                lines.append("</div>")

        elif section_key == "skills":
            if isinstance(data, list):
                lines.append(
                    '<p class="skills-list">'
                    + ", ".join(escape(str(s)) for s in data)
                    + "</p>"
                )
            elif isinstance(data, dict):
                for category, items in data.items():
                    lines.append('<p class="skill-category">')
                    lines.append(f"<strong>{escape(str(category))}:</strong> ")
                    if isinstance(items, list):
                        lines.append(", ".join(escape(str(s)) for s in items))
                    else:
                        lines.append(escape(str(items)))
                    lines.append("</p>")

        elif section_key == "projects":
            for proj in data if isinstance(data, list) else []:
                lines.append('<div class="entry">')
                proj_name = escape(proj.get("name", ""))
                if proj_name:
                    lines.append(f'<h3 class="entry-title">{proj_name}</h3>')

                description = proj.get("description", "")
                if description:
                    lines.append(f"<p>{escape(description)}</p>")

                technologies = proj.get("technologies", "")
                if technologies:
                    if isinstance(technologies, list):
                        technologies = ", ".join(str(t) for t in technologies)
                    lines.append(
                        f"<p><strong>Technologies:</strong> {escape(str(technologies))}</p>"
                    )

                url = proj.get("url", "")
                if url:
                    lines.append(
                        f"<p><strong>URL:</strong> {escape(str(url))}</p>"
                    )
                lines.append("</div>")

        elif section_key == "certifications":
            for cert in data if isinstance(data, list) else []:
                lines.append('<div class="entry">')
                cert_name = escape(cert.get("name", ""))
                issuer = escape(cert.get("issuer", ""))
                header = cert_name
                if issuer:
                    header = f"{cert_name} &mdash; {issuer}" if cert_name else issuer
                if header:
                    lines.append(f'<h3 class="entry-title">{header}</h3>')

                meta_parts = []
                cert_date = cert.get("date", "")
                if cert_date:
                    meta_parts.append(escape(str(cert_date)))
                cert_url = cert.get("url", "")
                if cert_url:
                    meta_parts.append(escape(str(cert_url)))
                if meta_parts:
                    lines.append(
                        '<p class="entry-meta">' + " | ".join(meta_parts) + "</p>"
                    )
                lines.append("</div>")

        lines.append("</div>")  # close .section
        return "\n".join(lines)

    def _get_default_css(self) -> str:
        """Return default professional CSS for resume PDF rendering."""
        return """
            @page {
                size: A4;
                margin: 1.5cm 2cm;
            }

            body {
                font-family: "Helvetica", "Arial", sans-serif;
                font-size: 10pt;
                line-height: 1.5;
                color: #1a1a2e;
                margin: 0;
                padding: 0;
            }

            .header {
                text-align: center;
                border-bottom: 2px solid #2c3e50;
                padding-bottom: 10px;
                margin-bottom: 16px;
            }

            .header h1 {
                font-size: 22pt;
                margin: 0 0 4px 0;
                color: #1a1a2e;
                letter-spacing: 1px;
            }

            .contact {
                font-size: 9pt;
                color: #555555;
                margin: 4px 0 0 0;
            }

            .section {
                margin-bottom: 14px;
            }

            .section-title {
                font-size: 13pt;
                color: #2c3e50;
                border-bottom: 1px solid #bdc3c7;
                padding-bottom: 3px;
                margin: 12px 0 8px 0;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }

            .entry {
                margin-bottom: 10px;
            }

            .entry-title {
                font-size: 11pt;
                margin: 0 0 2px 0;
                color: #1a1a2e;
            }

            .entry-meta {
                font-size: 9pt;
                color: #777777;
                margin: 0 0 4px 0;
            }

            p {
                margin: 3px 0;
            }

            ul {
                margin: 4px 0 4px 20px;
                padding: 0;
            }

            li {
                margin-bottom: 2px;
            }

            .skills-list {
                margin: 4px 0;
            }

            .skill-category {
                margin: 3px 0;
            }

            strong {
                color: #1a1a2e;
            }
        """
